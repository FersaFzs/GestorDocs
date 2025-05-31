from docx import Document
import re

def load_template(template_path):
    """Carga un archivo .docx y devuelve su contenido como texto y el objeto Document."""
    doc = Document(template_path)
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"
    return full_text, doc

def find_identifiers(text):
    """Busca identificadores como {NOMBRE} en el texto y devuelve una lista única."""
    identifiers = re.findall(r"\{.*?\}", text)
    return list(dict.fromkeys(identifiers))  # Elimina duplicados

def generate_document(template_doc, output_path, replacements):
    """Rellena la plantilla con los datos y guarda el nuevo documento."""
    doc = template_doc
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            # Create a list to hold the new runs after replacement
            new_runs_data = []
            current_text = paragraph.text
            
            # Find all occurrences of the key in the current paragraph text
            start_index = 0
            while True:
                idx = current_text.find(key, start_index)
                if idx == -1:
                    break
                
                # Add the text before the key
                if idx > start_index:
                    new_runs_data.append({'text': current_text[start_index:idx], 'bold': False, 'italic': False, 'underline': False})
                
                # Add the replacement value
                new_runs_data.append({'text': value, 'bold': False, 'italic': False, 'underline': False})
                
                start_index = idx + len(key)
            
            # Add any remaining text after the last key
            if start_index < len(current_text):
                new_runs_data.append({'text': current_text[start_index:], 'bold': False, 'italic': False, 'underline': False})

            # Clear existing runs and add new ones, preserving some basic formatting
            if key in paragraph.text:  # Only modify if the key was found
                paragraph.clear()
                for run_data in new_runs_data:
                    new_run = paragraph.add_run(run_data['text'])
                    new_run.bold = run_data['bold']
                    new_run.italic = run_data['italic']
                    new_run.underline = run_data['underline']
                # After processing one key, update the paragraph.text for the next key processing
                # This is a simplification; a more robust solution would re-evaluate the runs for each key
                # For now, we rely on the paragraph.text being updated implicitly for subsequent keys

    doc.save(output_path)