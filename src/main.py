from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QPushButton, QVBoxLayout, QWidget, QComboBox, QFileDialog, QMessageBox, QHBoxLayout, QScrollArea, QDialog
from PyQt5.QtGui import QColor
from PyQt5.QtCore import Qt, QUrl
from PyQt5.QtWidgets import QGraphicsDropShadowEffect
from PyQt5.QtWebEngineWidgets import QWebEngineView
from src.doc_processor import load_template, find_identifiers, generate_document
from docx import Document
import sys
import os
import win32print
import win32api
from docx2pdf import convert
import shutil # Added for file operations like copying

# Function to ensure templates directory and a default template exist
def setup_directories_and_templates():
    templates_dir = "templates"
    output_dir = "output"
    default_template_path = os.path.join(templates_dir, "example_template.docx")

    os.makedirs(templates_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    # Check if there are any .docx files in the templates directory
    if not any(f.endswith(".docx") for f in os.listdir(templates_dir)):
        doc = Document() # Create a new blank document
        doc.add_paragraph("Este es un documento de plantilla de ejemplo.")
        doc.add_paragraph("Complete la {INFORMACION} con sus datos.")
        doc.add_paragraph("Dirección de envío: {DIRECCION_ENVIO}")
        doc.save(default_template_path)
        print(f"Plantilla de ejemplo creada en: {default_template_path}")

class PreviewDialog(QDialog):
    def __init__(self, pdf_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Vista Previa del Documento")
        self.setGeometry(200, 200, 800, 800)
        self.setModal(True)

        self.layout = QVBoxLayout(self)

        self.web_view = QWebEngineView()
        self.web_view.setUrl(QUrl.fromLocalFile(pdf_path))
        self.layout.addWidget(self.web_view)

        # Buttons for confirmation or going back
        self.button_layout = QHBoxLayout()
        self.confirm_button = QPushButton("Confirmar")
        self.confirm_button.clicked.connect(self.accept)
        self.button_layout.addWidget(self.confirm_button)

        self.back_button = QPushButton("Volver")
        self.back_button.clicked.connect(self.reject)
        self.button_layout.addWidget(self.back_button)

        self.layout.addLayout(self.button_layout)

class GestorDocsWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("GestorDocs")
        self.setGeometry(100, 100, 600, 700)  # Tamaño inicial
        self.setMaximumHeight(700)  # Fijar altura máxima (ajustable según pantalla)

        # Widget y layout principal
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout(self.central_widget)
        self.layout.setSpacing(15)
        self.layout.setContentsMargins(30, 30, 30, 30)

        # Área desplazable para campos dinámicos
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.fields_container = QWidget()
        self.fields_layout = QVBoxLayout(self.fields_container)
        self.fields_layout.setSpacing(10)
        self.scroll_area.setWidget(self.fields_container)
        self.layout.addWidget(self.scroll_area)
        self.fields = {}

        # Contenedor para botones (fuera del scroll)
        self.buttons_container = QWidget()
        self.buttons_layout = QVBoxLayout(self.buttons_container)
        self.buttons_layout.setSpacing(15)

        # Menú de plantillas
        self.template_combo = QComboBox()
        self.template_combo.addItem("Seleccionar plantilla")
        for file in os.listdir("templates"):
            if file.endswith(".docx"):
                self.template_combo.addItem(file)
        self.template_combo.currentTextChanged.connect(self.load_template_fields)
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(10)
        shadow.setXOffset(3)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 80))
        self.template_combo.setGraphicsEffect(shadow)
        self.buttons_layout.addWidget(self.template_combo)

        # Subcontenedor para los tres botones
        self.action_buttons_container = QWidget()
        self.action_buttons_layout = QHBoxLayout(self.action_buttons_container)
        self.action_buttons_layout.setSpacing(20)

        self.save_btn = QPushButton("Guardar como")
        self.save_btn.clicked.connect(self.save_doc)
        self.save_btn.setEnabled(False)
        self.action_buttons_layout.addWidget(self.save_btn)
        self.action_buttons_layout.addStretch()

        self.generate_btn = QPushButton("Generar documento")
        self.generate_btn.clicked.connect(self.generate_doc)
        self.generate_btn.setEnabled(False)
        self.action_buttons_layout.addWidget(self.generate_btn)
        self.action_buttons_layout.addStretch()

        self.print_btn = QPushButton("Imprimir documento")
        self.print_btn.clicked.connect(self.print_doc)
        self.print_btn.setEnabled(False)
        self.action_buttons_layout.addWidget(self.print_btn)

        # Import Template Button
        self.import_template_btn = QPushButton("Importar Plantilla")
        self.import_template_btn.clicked.connect(self.import_template)
        self.action_buttons_layout.addWidget(self.import_template_btn)

        self.buttons_layout.addWidget(self.action_buttons_container)
        self.layout.addWidget(self.buttons_container)

    def load_template_fields(self, template_name):
        for field in self.fields.values():
            field.deleteLater()
        self.fields.clear()
        
        while self.fields_layout.count():
            child = self.fields_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        if template_name == "Seleccionar plantilla":
            self.generate_btn.setEnabled(False)
            self.save_btn.setEnabled(False)
            self.print_btn.setEnabled(False)
            return

        try:
            template_path = os.path.join("templates", template_name)
            text, self.doc = load_template(template_path)
            identifiers = find_identifiers(text)
            for identifier in identifiers:
                label = QLabel(identifier.strip("{}") + ":")
                self.fields_layout.addWidget(label)
                entry = QLineEdit()
                shadow = QGraphicsDropShadowEffect(self)
                shadow.setBlurRadius(10)
                shadow.setXOffset(3)
                shadow.setYOffset(3)
                shadow.setColor(QColor(0, 0, 0, 80))
                entry.setGraphicsEffect(shadow)
                self.fields[identifier] = entry
                self.fields_layout.addWidget(entry)
            self.generate_btn.setEnabled(True)
        except Exception as e:
            QMessageBox.critical(self, "Error al cargar plantilla", f"No se pudo cargar la plantilla: {template_name}\nMotivo: {str(e)}\nAsegúrate de que el archivo no esté corrupto o que Microsoft Word esté instalado si es necesario.")
            self.generate_btn.setEnabled(False)
            self.save_btn.setEnabled(False)
            self.print_btn.setEnabled(False)
            # Clear fields if template loading fails
            for field in self.fields.values():
                field.deleteLater()
            self.fields.clear()
            while self.fields_layout.count():
                child = self.fields_layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()

    def generate_doc(self):
        if not self.fields:
            return
        replacements = {key: entry.text() for key, entry in self.fields.items()}
        self.temp_docx_path = os.path.join("output", "temp_documento.docx")
        generate_document(self.doc, self.temp_docx_path, replacements)
        print(f"Documento temporal generado en {self.temp_docx_path}")

        # Convert DOCX to PDF for preview
        self.temp_pdf_path = os.path.join("output", "temp_documento.pdf")
        try:
            convert(self.temp_docx_path, self.temp_pdf_path)
            print(f"Documento temporal convertido a PDF para vista previa en {self.temp_pdf_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error de Conversión", f"No se pudo generar la vista previa en PDF.\nMotivo: {str(e)}\nAsegúrate de tener Microsoft Word instalado.")
            return

        # Show preview dialog
        preview_dialog = PreviewDialog(self.temp_pdf_path, self)
        if preview_dialog.exec_() == QDialog.Accepted:
            # User confirmed, enable save/print buttons
            self.save_btn.setEnabled(True)
            self.print_btn.setEnabled(True)
            self.last_output_path = self.temp_docx_path # Set last_output_path for saving/printing
        else:
            # User went back, disable save/print buttons and clean up temporary files
            self.save_btn.setEnabled(False)
            self.print_btn.setEnabled(False)
            # Clear last_output_path if user cancels preview
            if hasattr(self, 'last_output_path'):
                del self.last_output_path
            if os.path.exists(self.temp_docx_path):
                os.remove(self.temp_docx_path)
            if os.path.exists(self.temp_pdf_path):
                os.remove(self.temp_pdf_path)
            print("Vista previa cancelada. Archivos temporales eliminados.")

    def save_doc(self):
        if not hasattr(self, 'temp_docx_path'):
            return
        output_path, _ = QFileDialog.getSaveFileName(self, "Guardar documento", "output/", "Documentos (*.docx *.pdf)")
        if output_path:
            try:
                if output_path.endswith(".pdf"):
                    convert(self.temp_docx_path, output_path)
                    print(f"Guardado como PDF en {output_path}")
                else:
                    os.rename(self.temp_docx_path, output_path)
                    print(f"Guardado como DOCX en {output_path}")
                self.last_output_path = output_path
            except Exception as e:
                error_message = f"No se pudo guardar como PDF.\nMotivo: {str(e)}"
                if "No module named 'win32com'" in str(e):
                    error_message += "\nAsegúrate de tener pywin32 instalado (pip install pywin32)."
                elif "Microsoft Word is not installed" in str(e) or "Word application cannot be found" in str(e):
                    error_message += "\nAsegúrate de tener Microsoft Word instalado y activado."
                else:
                    error_message += "\nVerifica que el archivo no esté abierto en otra aplicación y que tienes permisos de escritura."
                QMessageBox.critical(self, "Error", error_message)
                return
            finally:
                if os.path.exists(self.temp_docx_path):
                    os.remove(self.temp_docx_path)
                    print(f"Archivo temporal {self.temp_docx_path} eliminado")

    def print_doc(self):
        if not hasattr(self, 'last_output_path') or not os.path.exists(self.last_output_path):
            QMessageBox.warning(self, "Advertencia", "No hay un documento guardado para imprimir.")
            return

        try:
            win32api.ShellExecute(0, "print", self.last_output_path, None, ".", 0)
            QMessageBox.information(self, "Éxito", f"Documento enviado a la impresora: {os.path.basename(self.last_output_path)}")
            print("Enviado a la impresora")
        except Exception as e:
            error_message = f"No se pudo enviar a la impresora.\nMotivo: {str(e)}"
            if "El sistema no puede encontrar el archivo especificado" in str(e):
                error_message += "\nAsegúrate de que el documento exista y la ruta sea correcta."
            elif "No se encontró ninguna impresora" in str(e) or "No printers are installed" in str(e):
                error_message += "\nAsegúrate de tener una impresora instalada y configurada."
            else:
                error_message += "\nVerifica tu configuración de impresora o el estado del archivo."
            QMessageBox.critical(self, "Error de impresión", error_message)

    def import_template(self):
        # Show reminder about template pattern
        QMessageBox.information(self,
                                "Importar Plantilla",
                                "Recuerda: la plantilla debe tener los espacios a rellenar con la palabra clave entre llaves, ej: {NOMBRE_CAMPO}.")

        # Open file dialog to select .docx file
        file_path, _ = QFileDialog.getOpenFileName(self, "Seleccionar Plantilla", "", "Documentos Word (*.docx)")

        if file_path:
            try:
                template_name = os.path.basename(file_path)
                destination_path = os.path.join("templates", template_name)
                
                # Ensure templates directory exists
                os.makedirs("templates", exist_ok=True)

                # Copy the file to the templates directory
                shutil.copyfile(file_path, destination_path)
                QMessageBox.information(self, "Importación Exitosa", f"Plantilla '{template_name}' importada con éxito.")
                self.refresh_template_combo()
            except Exception as e:
                QMessageBox.critical(self, "Error de Importación", f"No se pudo importar la plantilla.\nMotivo: {str(e)}")
    
    def refresh_template_combo(self):
        self.template_combo.clear()
        self.template_combo.addItem("Seleccionar plantilla")
        for file in os.listdir("templates"):
            if file.endswith(".docx"):
                self.template_combo.addItem(file)

if __name__ == "__main__":
    # Ensure directories and default templates are set up before running the app
    setup_directories_and_templates()

    app = QApplication(sys.argv)
    with open("src/styles.qss", "r") as style_file:
        app.setStyleSheet(style_file.read())
    window = GestorDocsWindow()
    window.show()
    sys.exit(app.exec_())