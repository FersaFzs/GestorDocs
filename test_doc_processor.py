import pytest
from docx import Document # For creating new documents
from docx.document import Document as DocxDocument  # For type checking against the Document class
import os
from src.doc_processor import load_template, find_identifiers, generate_document

# Define a path for a dummy template and output file
DUMMY_TEMPLATE_PATH = "templates/dummy_template.docx"
TEST_OUTPUT_PATH = "output/test_output.docx"

@pytest.fixture(scope="module", autouse=True)
def setup_dummy_template():
    # Create a dummy .docx file for testing
    doc = Document()
    doc.add_paragraph("This is a test document with {KEYWORD1} and {KEYWORD2} and {KEYWORD1} repeated.")
    doc.add_paragraph("Another paragraph with {ANOTHER_KEYWORD}.")
    
    # Ensure the templates directory exists
    os.makedirs(os.path.dirname(DUMMY_TEMPLATE_PATH), exist_ok=True)
    doc.save(DUMMY_TEMPLATE_PATH)
    
    # Ensure the output directory exists
    os.makedirs(os.path.dirname(TEST_OUTPUT_PATH), exist_ok=True)
    
    yield
    
    # Clean up the dummy file and output file after tests
    if os.path.exists(DUMMY_TEMPLATE_PATH):
        os.remove(DUMMY_TEMPLATE_PATH)
    if os.path.exists(TEST_OUTPUT_PATH):
        os.remove(TEST_OUTPUT_PATH)

def test_load_template():
    text, doc = load_template(DUMMY_TEMPLATE_PATH)
    assert "This is a test document" in text
    assert isinstance(doc, DocxDocument) # Use the aliased Document for type checking

def test_find_identifiers():
    text, _ = load_template(DUMMY_TEMPLATE_PATH)
    identifiers = find_identifiers(text)
    assert "{KEYWORD1}" in identifiers
    assert "{KEYWORD2}" in identifiers
    assert "{ANOTHER_KEYWORD}" in identifiers
    assert len(identifiers) == 3  # Ensure no duplicates and correct count

def test_generate_document():
    text, doc = load_template(DUMMY_TEMPLATE_PATH)
    replacements = {
        "{KEYWORD1}": "VALUE1",
        "{KEYWORD2}": "VALUE2",
        "{ANOTHER_KEYWORD}": "VALUE3"
    }
    generate_document(doc, TEST_OUTPUT_PATH, replacements)
    
    # Verify the output document content (this requires re-loading and checking)
    output_doc = Document(TEST_OUTPUT_PATH)
    output_text = ""
    for paragraph in output_doc.paragraphs:
        output_text += paragraph.text
    
    assert "VALUE1" in output_text
    assert "VALUE2" in output_text
    assert "VALUE3" in output_text
    assert "{KEYWORD1}" not in output_text # Ensure all occurrences are replaced
    assert "{KEYWORD2}" not in output_text
    assert "{ANOTHER_KEYWORD}" not in output_text 